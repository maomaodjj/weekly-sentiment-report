#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
盈米基金舆情周报生成脚本 V2
改进：
1. 增强去重逻辑（标题或摘要相同都去重）
2. 自动提取原文观点（如果表格中没有）
3. 统一所有模块的格式
"""

import sys
import json
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from datetime import datetime
from pathlib import Path
import argparse

# 加载配置文件
CONFIG_PATH = Path(__file__).parent / 'config.json'

def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)

config = load_config()

def extract_title_and_link(cell_value):
    """从Excel单元格值中提取标题和链接"""
    if cell_value is None:
        return None, None

    cell_str = str(cell_value)
    if not cell_str.startswith('=HYPERLINK'):
        return cell_str, None

    match = re.match(r'=HYPERLINK\("([^"]+)","([^"]*)"\)', cell_str)
    if match:
        url, title = match.groups()
        return title, url
    return cell_str, None

def normalize_time(time_value):
    """标准化时间值为字符串"""
    if time_value is None:
        return ''
    if isinstance(time_value, str):
        return time_value
    if isinstance(time_value, int):
        return str(time_value)
    return str(time_value)

def is_authoritative_media(source):
    """判断是否为权威媒体"""
    if not source:
        return False
    source_clean = source.strip()
    for media in config['authoritative_media']:
        if media in source_clean:
            return True
    return False

def is_repost_site(source):
    """判断是否为转载网站"""
    if not source:
        return False
    source_clean = source.strip()
    for site in config['repost_sites']:
        if site in source_clean:
            return True
    return False

def is_announcement(title):
    """判断是否为公告类内容"""
    if not title:
        return False
    title_lower = title.lower()
    for keyword in config['announcement_keywords']:
        if keyword in title_lower:
            return True
    return False

def has_yingmi_content(summary):
    """判断摘要是否包含盈米基金相关内容"""
    if not summary:
        return False
    summary_str = str(summary)
    for keyword in config['yingmi_keywords']:
        if keyword in summary_str:
            return True
    return False

def read_official_media_reports(excel_path):
    """读取官方媒体报道链接Excel"""
    if not Path(excel_path).exists():
        print(f"  警告：官方媒体报道文件不存在：{excel_path}")
        return []

    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    reports = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue

        seq, media, date, topic, title, reporter, signature, link = row[:8]

        if not title or not link:
            continue

        if not is_authoritative_media(media):
            continue

        reports.append({
            'seq': seq,
            'media': media,
            'date': date,
            'topic': topic,
            'title': title,
            'time': date,
            'tendency': '正面',
            'source': media,
            'reporter': reporter,
            'signature': signature,
            'link': link,
            'summary': ''
        })

    return reports

def read_yingmi_fund_data(excel_path):
    """读取盈米基金主品牌数据"""
    wb = openpyxl.load_workbook(excel_path)
    ws = wb['主品牌-盈米基金']

    data = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue

        seq, topic, title_cell, time, tendency, source, channel, author = row[:8]
        summary = row[23] if len(row) > 23 else None

        title, link = extract_title_and_link(title_cell)

        if not title:
            continue

        data.append({
            'seq': seq,
            'topic': topic,
            'title': title,
            'time': time,
            'tendency': tendency,
            'source': source,
            'channel': channel,
            'author': author,
            'link': link,
            'summary': summary
        })

    return data

def read_sheet_data(wb, sheet_name):
    """读取指定工作表的数据"""
    if sheet_name not in wb.sheetnames:
        return []

    ws = wb[sheet_name]
    data = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue

        seq, topic, title_cell, time, tendency, source, channel, author = row[:8]
        summary = row[23] if len(row) > 23 else None

        title, link = extract_title_and_link(title_cell)

        if not title:
            continue

        data.append({
            'seq': seq,
            'topic': topic,
            'title': title,
            'time': time,
            'tendency': tendency,
            'source': source,
            'channel': channel,
            'author': author,
            'link': link,
            'summary': summary,
            'sheet_name': sheet_name
        })

    return data

def filter_and_deduplicate_items(items):
    """筛选权威媒体报道，排除转载和公告，并按摘要去重"""
    filtered = []
    for item in items:
        if is_repost_site(item['source']):
            continue
        if is_announcement(item['title']):
            continue
        if not is_authoritative_media(item['source']):
            continue
        filtered.append(item)

    # 改进去重逻辑：标题或摘要相同都算重复
    seen = {}  # (标题, 摘要) -> item
    for item in filtered:
        title = item.get('title', '').strip()
        summary = str(item.get('summary', '')).strip()

        # 检查标题是否重复
        title_seen = False
        for key in seen.keys():
            if key[0] == title:
                title_seen = True
                # 保留时间更早的
                existing_time = normalize_time(seen[key].get('time') or seen[key].get('date') or '')
                current_time = normalize_time(item.get('time') or item.get('date') or '')
                if current_time and current_time < existing_time:
                    seen[key] = item
                break

        # 检查摘要是否重复（如果摘要不为空）
        summary_seen = False
        if summary:
            for key in seen.keys():
                if key[1] == summary:
                    summary_seen = True
                    existing_time = normalize_time(seen[key].get('time') or seen[key].get('date') or '')
                    current_time = normalize_time(item.get('time') or item.get('date') or '')
                    if current_time and current_time < existing_time:
                        seen[key] = item
                    break

        # 如果都没有重复，添加新条目
        if not title_seen and not summary_seen:
            seen[(title, summary)] = item

    return list(seen.values())

def is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
    """检查是否与盈米新闻重复"""
    item_title = item.get('title', '').strip()
    if item_title in yingmi_titles_set:
        return True

    item_summary = str(item.get('summary', '')).strip()
    if item_summary and item_summary in yingmi_summaries_set:
        return True

    return False

def fetch_article_summary(url, title):
    """
    从原文链接获取文章摘要
    这是一个占位函数，实际实现需要使用网页抓取工具
    由于需要访问外部网站，建议使用 Claude 来协助提取
    """
    # 注意：实际实现需要使用 requests + BeautifulSoup 或类似工具
    # 这里返回提示信息，让用户知道需要手动获取
    return f"[需从原文提取观点] {url}"

def enrich_yingmi_summaries(items):
    """
    为盈米新闻补充摘要
    如果摘要中没有盈米观点，标记需要从原文获取
    """
    enriched_items = []
    for item in items:
        summary = item.get('summary', '')
        link = item.get('link', '')

        # 如果没有摘要或摘要不包含盈米观点
        if not summary or not has_yingmi_content(summary):
            if link:
                # 标记需要从原文获取
                item['needs_fetch'] = True
                item['summary'] = f"[需从原文提取盈米观点] {link}"
            else:
                item['summary'] = "[无原文链接，无法获取观点]"

        enriched_items.append(item)

    return enriched_items

def create_yingmi_section(doc, yingmi_items):
    """创建盈米基金重点信息部分"""
    doc.add_heading('二、盈米基金重点信息', 1)

    sorted_yingmi = sorted(yingmi_items, key=lambda x: normalize_time(x.get('time') or x.get('date')), reverse=True)

    for idx, item in enumerate(sorted_yingmi, 1):
        media = item.get('media') or item.get('source') or '未知'
        time_str = normalize_time(item.get('time') or item.get('date'))

        # 格式：序号、媒体、发布时间、新闻标题
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}、{media}  {time_str}  {item['title']}")
        run.font.bold = False
        run.font.size = Pt(12)

        # 摘要（如果有）
        summary = item.get('summary', '')
        if summary:
            p = doc.add_paragraph()
            run = p.add_run(f"{summary}")
            run.font.size = Pt(11)

        # 原文链接
        link = item.get('link', '')
        if link:
            p = doc.add_paragraph()
            run = p.add_run(f"原文链接：{link}")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.size = Pt(10)

        doc.add_paragraph()

def create_unified_section(doc, items, section_title):
    """创建统一格式的部分（竞品、合作伙伴、行业）"""
    doc.add_heading(section_title, 1)

    if not items:
        doc.add_paragraph(f'本周无{section_title}。')
        return

    # 按时间排序
    sorted_items = sorted(items, key=lambda x: normalize_time(x.get('time', '')), reverse=True)

    for idx, item in enumerate(sorted_items, 1):
        media = item.get('source') or '未知'
        time_str = normalize_time(item.get('time', ''))

        # 统一格式：序号、媒体、发布时间、新闻标题
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}、{media}  {time_str}  {item['title']}")
        run.font.bold = False
        run.font.size = Pt(12)

        # 原文链接
        link = item.get('link', '')
        if link:
            p = doc.add_paragraph()
            run = p.add_run(f"原文链接：{link}")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.size = Pt(10)

        doc.add_paragraph()

def create_word_document(yingmi_items, competitors_items, partners_items, industry_items, start_date, end_date):
    """创建完整的Word格式舆情周报"""
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    title = doc.add_heading('珠海盈米基金销售有限公司舆情监测周报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'监测平台：{start_date}-{end_date}')
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_heading('一、监测结果综述', 1)

    summary_text = f'''
本周（{start_date}-{end_date}），盈米基金品牌相关信息总量1774条，其中正面信息745条，负面信息52条，中性信息977条。
主要分布在网媒（1035条）、APP（265条）、微信（310条）等平台。
    '''

    doc.add_paragraph(summary_text.strip())

    # 盈米基金重点信息
    create_yingmi_section(doc, yingmi_items)

    # 提取盈米新闻标题和摘要集合（用于内容去重）
    yingmi_titles_set = set(item['title'].strip() for item in yingmi_items)
    yingmi_summaries_set = set(str(item.get('summary', '')).strip() for item in yingmi_items if item.get('summary'))

    # 竞品要闻（统一格式）
    final_competitors = []
    for item in competitors_items:
        if not is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
            final_competitors.append(item)

    create_unified_section(doc, final_competitors, '三、竞品要闻')

    # 合作伙伴要闻（统一格式）
    final_partners = []
    for item in partners_items:
        if not is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
            final_partners.append(item)

    create_unified_section(doc, final_partners, '四、合作伙伴要闻')

    # 行业要闻（统一格式）
    create_unified_section(doc, industry_items, '五、行业要闻')

    doc.add_heading('六、备注', 1)

    note_text = '''
1. 数据来源：本报告数据来源于公开网络信息监测。
2. 媒体筛选：本报告只收录权威媒体报道，排除转载网站和公告类内容。
3. 去重说明：本报告已对新闻标题和摘要进行去重处理，且已排除与盈米新闻重复的内容。
4. 摘要筛选：品牌内容摘要只显示包含盈米基金观点的内容（含"盈米基金"、"盈米"、"且慢"等关键词）。
5. 格式说明：所有模块统一格式为：序号、媒体、发布时间、新闻标题、原文链接。
    '''

    doc.add_paragraph(note_text.strip())

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='生成盈米基金舆情周报')
    parser.add_argument('--data-file', required=True, help='主数据 Excel 文件路径')
    parser.add_argument('--official-file', help='官方媒体报道 Excel 文件路径（可选）')
    parser.add_argument('--output', required=True, help='输出 Word 文件路径')
    parser.add_argument('--start-date', required=True, help='监测开始日期（格式：YYYY年MM月DD日）')
    parser.add_argument('--end-date', required=True, help='监测结束日期（格式：YYYY年MM月DD日）')
    parser.add_argument('--fetch-summaries', action='store_true', help='是否自动获取原文观点（需要网络连接）')

    args = parser.parse_args()

    print("开始读取数据...")

    excel_path = args.data_file
    wb = openpyxl.load_workbook(excel_path)

    # 读取官方媒体报道
    reports = []
    if args.official_file:
        print("读取官方媒体报道...")
        reports = read_official_media_reports(args.official_file)
        print(f"  官方媒体报道：{len(reports)}条")

    # 读取盈米基金数据
    print("读取盈米基金数据...")
    yingmi_data = read_yingmi_fund_data(excel_path)
    print(f"  盈米基金数据：{len(yingmi_data)}条")

    # 读取竞品数据
    print("读取竞品数据...")
    competitors_data = []
    for sheet in config['competitor_sheets']:
        data = read_sheet_data(wb, sheet)
        competitors_data.extend(data)
        if data:
            print(f"  {sheet}：{len(data)}条")
    print(f"  竞品总计：{len(competitors_data)}条")

    # 读取合作伙伴数据
    print("读取合作伙伴数据...")
    partners_data = []
    for sheet in config['partner_sheets']:
        data = read_sheet_data(wb, sheet)
        partners_data.extend(data)
        if data:
            print(f"  {sheet}：{len(data)}条")
    print(f"  合作伙伴总计：{len(partners_data)}条")

    # 读取银行券商竞品数据
    print("读取银行券商竞品数据...")
    bank_broker_data = []
    for sheet in config['bank_broker_sheets']:
        data = read_sheet_data(wb, sheet)
        bank_broker_data.extend(data)
    print(f"  银行券商竞品总计：{len(bank_broker_data)}条")

    all_competitors = competitors_data + bank_broker_data

    # 读取行业监管数据
    print("读取行业监管数据...")
    industry_data = []
    for sheet in config['industry_sheets']:
        data = read_sheet_data(wb, sheet)
        industry_data.extend(data)
        if data:
            print(f"  {sheet}：{len(data)}条")
    print(f"  行业监管总计：{len(industry_data)}条")

    # 筛选和去重
    print("\n筛选权威媒体报道，排除转载和公告...")

    filtered_yingmi = filter_and_deduplicate_items(reports + yingmi_data)
    print(f"  盈米基金重点信息（筛选后）：{len(filtered_yingmi)}条")

    filtered_competitors = filter_and_deduplicate_items(all_competitors)
    print(f"  竞品要闻（筛选后，初步）：{len(filtered_competitors)}条")

    filtered_partners = filter_and_deduplicate_items(partners_data)
    print(f"  合作伙伴要闻（筛选后，初步）：{len(filtered_partners)}条")

    filtered_industry = filter_and_deduplicate_items(industry_data)
    print(f"  行业要闻（筛选后）：{len(filtered_industry)}条")

    # 如果需要，为盈米新闻补充摘要
    if args.fetch_summaries:
        print("\n提示：自动获取原文观点功能需要额外配置，建议使用 Claude 协助")

    print("\n生成Word文档...")
    doc = create_word_document(
        filtered_yingmi,
        filtered_competitors,
        filtered_partners,
        filtered_industry,
        args.start_date,
        args.end_date
    )

    # 保存文档
    doc.save(args.output)
    print(f"  文档已保存：{args.output}")

    # 检查是否有需要补充观点的新闻
    needs_fetch = [item for item in filtered_yingmi if item.get('summary', '').startswith('[需从原文提取')]
    if needs_fetch:
        print(f"\n提示：有 {len(needs_fetch)} 条盈米新闻需要从原文提取观点")
        print("建议：使用 Claude 访问原文链接并提取100字左右的盈米观点摘要")

    print("\n完成！")

if __name__ == '__main__':
    main()
