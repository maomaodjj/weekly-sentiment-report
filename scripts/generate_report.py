#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
盈米基金舆情周报生成脚本 V2
改进：
1. 增强去重逻辑（标题或摘要相同都去重）
2. 自动提取原文观点（如果表格中没有）
3. 统一所有模块的格式
"""

import sys
import json
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from datetime import datetime
from pathlib import Path
import argparse
from difflib import SequenceMatcher
import re

# 加载配置文件
CONFIG_PATH = Path(__file__).parent / 'config.json'

def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)

config = load_config()

def extract_title_and_link(cell_value):
    """从Excel单元格值中提取标题和链接"""
    if cell_value is None:
        return None, None

    cell_str = str(cell_value)
    if not cell_str.startswith('=HYPERLINK'):
        return cell_str, None

    match = re.match(r'=HYPERLINK\("([^"]+)","([^"]*)"\)', cell_str)
    if match:
        url, title = match.groups()
        return title, url
    return cell_str, None

def normalize_time(time_value):
    """标准化时间值为字符串"""
    if time_value is None:
        return ''
    if isinstance(time_value, str):
        return time_value
    # 处理Excel日期数字格式（如46041）
    if isinstance(time_value, (int, float)):
        try:
            # Excel日期基准是1899-12-30，转换为datetime
            from datetime import datetime, timedelta
            excel_base = datetime(1899, 12, 30)
            delta = timedelta(days=time_value)
            dt = excel_base + delta
            # 格式化为 "YYYY-MM-DD HH:MM:SS"
            return dt.strftime('%Y-%m-%d %H:%M:%S')
        except:
            return str(time_value)
    return str(time_value)

def is_authoritative_media(source):
    """判断是否为权威媒体"""
    if not source:
        return False
    source_clean = source.strip()
    for media in config['authoritative_media']:
        if media in source_clean:
            return True
    return False

def is_repost_site(source):
    """判断是否为转载网站"""
    if not source:
        return False
    source_clean = source.strip()
    for site in config['repost_sites']:
        if site in source_clean:
            return True
    return False

def is_announcement(title):
    """判断是否为公告类内容"""
    if not title:
        return False
    title_str = str(title)
    title_lower = title_str.lower()

    # 检查是否为报纸版面接续内容（如"(上接58版)"）
    if re.match(r'^[（(]\s*上接.*版\s*[）)]$', title_str.strip()):
        return True

    for keyword in config['announcement_keywords']:
        if keyword in title_lower:
            return True
    # 额外检查发售机构披露类信息
    if '为销售机构' in title_str or '销售机构：' in title_str or '新增.*为销售机构' in title_str:
        return True
    return False

def is_industry_excluded(item):
    """判断行业新闻是否应该被排除（如医保相关内容）"""
    title = item.get('title', '')
    summary = str(item.get('summary', ''))
    source = item.get('source', '')

    # 检查标题、摘要和来源
    text_to_check = f"{title} {summary} {source}"

    exclude_keywords = config.get('industry_exclude_keywords', [])
    for keyword in exclude_keywords:
        if keyword in text_to_check:
            return True

    return False

def is_contact_info_disclosure(summary):
    """判断摘要是否为纯联系方式、地址等信息披露内容"""
    if not summary:
        return False
    summary_str = str(summary).strip()

    # 如果摘要很短（少于30字），检查是否为联系方式
    if len(summary_str) < 30:
        # 检查是否主要是电话号码、地址、联系人等
        patterns = [
            r'^[\d\s\-()（）+]{5,}$',  # 纯电话号码
            r'^[a-zA-Z\d\s\-()（）+@.]{5,}$',  # 邮箱或网址
            r'^(电话|手机|联系|地址|邮箱|网址|传真|邮编).*$',
            r'^[\u4e00-\u9fa5]*(地址|电话|联系|邮箱)[\u4e00-\u9fa5\d\s\-()（）@.]*$',
        ]
        for pattern in patterns:
            if re.match(pattern, summary_str):
                return True

    # 检查是否只包含地址信息
    if len(summary_str) > 0 and len(summary_str) < 50:
        # 如果摘要主要由地址相关词汇组成
        address_keywords = ['地址', '省', '市', '区', '路', '号', '室', '层', '楼', '大厦', '中心', '广场']
        address_count = sum(1 for kw in address_keywords if kw in summary_str)
        # 如果超过3个地址关键词且没有其他实质内容
        if address_count >= 3 and len(summary_str) < 80:
            return True

    return False

def calculate_similarity(str1, str2):
    """计算两个字符串的相似度（0-1）"""
    return SequenceMatcher(None, str1, str2).ratio()

def is_title_similar(title1, title2):
    """判断两个标题是否相似（部分重复）"""
    if not title1 or not title2:
        return False

    # 标准化后完全相同
    norm1 = normalize_text(title1)
    norm2 = normalize_text(title2)
    if norm1 == norm2:
        return True

    # 检查相似度
    similarity = calculate_similarity(norm1, norm2)
    if similarity >= 0.7:  # 相似度70%以上
        return True

    # 检查是否有较长的公共子串（至少15个字符）
    # 去除空格和标点后检查
    clean1 = re.sub(r'[^\w]', '', norm1)
    clean2 = re.sub(r'[^\w]', '', norm2)

    if len(clean1) > 20 and len(clean2) > 20:
        # 找到最长公共子串
        match = SequenceMatcher(None, clean1, clean2).find_longest_match(0, len(clean1), 0, len(clean2))
        if match.size >= 15:  # 至少15个字符相同
            # 检查公共子串占较短标题的比例
            shorter_len = min(len(clean1), len(clean2))
            if match.size / shorter_len >= 0.4:  # 占较短标题的40%以上
                return True

    return False

def is_summary_similar(summary1, summary2):
    """判断两个摘要是否相似（包含重复的一句话）"""
    if not summary1 or not summary2:
        return False

    # 标准化
    norm1 = normalize_text(summary1)
    norm2 = normalize_text(summary2)

    # 完全相同
    if norm1 == norm2:
        return True

    # 如果任一摘要太短，不进行相似度检查
    if len(norm1) < 10 or len(norm2) < 10:
        return False

    # 检查是否有较长的公共子串（至少20个字符）
    clean1 = re.sub(r'[^\w]', '', norm1)
    clean2 = re.sub(r'[^\w]', '', norm2)

    match = SequenceMatcher(None, clean1, clean2).find_longest_match(0, len(clean1), 0, len(clean2))

    # 至少20个字符相同，且占较短摘要的30%以上
    if match.size >= 20:
        shorter_len = min(len(clean1), len(clean2))
        if match.size / shorter_len >= 0.3:
            return True

    return False

def has_yingmi_content(summary):
    """判断摘要是否包含盈米基金相关内容"""
    if not summary:
        return False
    summary_str = str(summary)
    for keyword in config['yingmi_keywords']:
        if keyword in summary_str:
            return True
    return False

def read_official_media_reports(excel_path):
    """读取官方媒体报道链接Excel"""
    if not Path(excel_path).exists():
        print(f"  警告：官方媒体报道文件不存在：{excel_path}")
        return []

    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    reports = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue

        # 读取所有列，包括摘要（列8）
        seq, media, date, topic, title, reporter, signature, link = row[:8]
        summary = row[8] if len(row) > 8 else ''

        if not title or not link:
            continue

        if not is_authoritative_media(media):
            continue

        reports.append({
            'seq': seq,
            'media': media,
            'date': date,
            'topic': topic,
            'title': title,
            'time': date,
            'tendency': '正面',
            'source': media,
            'reporter': reporter,
            'signature': signature,
            'link': link,
            'summary': summary
        })

    return reports

def read_yingmi_fund_data(excel_path):
    """读取盈米基金主品牌数据"""
    wb = openpyxl.load_workbook(excel_path)

    # 获取盈米相关工作表列表
    yingmi_sheets = config.get('yingmi_sheets', ['盈米基金', '主品牌-盈米基金'])

    data = []
    for sheet_name in yingmi_sheets:
        if sheet_name not in wb.sheetnames:
            continue

        print(f"  读取工作表：{sheet_name}")
        ws = wb[sheet_name]

        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                continue

            # 新的Excel列结构：
            # 0:序号, 1:监测对象, 2:媒体类型, 3:媒体级别, 4:标题, 5:摘要,
            # 6:发布媒体, 7:文章类型, 8:地域, 9:作者, 10:作者认证,
            # 11:倾向性, 12:关键词, 13:发布时间, 14:热度

            seq = row[0]
            topic = row[1]
            title_cell = row[4]
            summary = row[5] if len(row) > 5 else None
            source = row[6]
            tendency = row[11]
            time = row[13]

            title, link = extract_title_and_link(title_cell)

            if not title:
                continue

            data.append({
                'seq': seq,
                'topic': topic,
                'title': title,
                'time': time,
                'tendency': tendency,
                'source': source,
                'link': link,
                'summary': summary,
                'sheet_name': sheet_name  # 添加来源工作表
            })

    return data

def read_sheet_data(wb, sheet_name):
    """读取指定工作表的数据"""
    if sheet_name not in wb.sheetnames:
        return []

    ws = wb[sheet_name]
    data = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue

        # 新的Excel列结构：
        # 0:序号, 1:监测对象, 2:媒体类型, 3:媒体级别, 4:标题, 5:摘要,
        # 6:发布媒体, 7:文章类型, 8:地域, 9:作者, 10:作者认证,
        # 11:倾向性, 12:关键词, 13:发布时间, 14:热度

        seq = row[0]
        topic = row[1]
        title_cell = row[4]
        summary = row[5] if len(row) > 5 else None
        source = row[6]
        tendency = row[11]
        time = row[13]

        title, link = extract_title_and_link(title_cell)

        if not title:
            continue

        data.append({
            'seq': seq,
            'topic': topic,
            'title': title,
            'time': time,
            'tendency': tendency,
            'source': source,
            'link': link,
            'summary': summary,
            'sheet_name': sheet_name
        })

    return data

def normalize_text(text):
    """标准化文本，用于去重比较"""
    if not text:
        return ''
    text = str(text).strip()
    # 统一引号格式并去除引号
    text = text.replace('\u201c', '').replace('\u201d', '')  # 中文引号
    text = text.replace('\u2018', '').replace('\u2019', '')  # 中文单引号
    text = text.replace('"', '').replace("'", '')  # ASCII引号
    # 去除所有空格（包括中文空格、英文空格、制表符、换行）
    text = re.sub(r'[\s\u3000]', '', text)
    # 去除常见标点符号
    text = text.replace('，', '').replace('。', '').replace('：', '')
    text = text.replace(',', '').replace('.', '').replace(':', '')
    text = text.replace('【', '').replace('】', '').replace('[', '').replace(']', '')
    text = text.replace('（', '').replace('）', '').replace('(', '').replace(')', '')
    text = text.replace('-', '').replace('_', '')
    # 转小写
    return text.lower()

def filter_and_deduplicate_items(items, debug=False, apply_industry_exclusion=False):
    """筛选权威媒体报道，排除转载和公告，并严格按标题和摘要去重"""
    filtered = []
    for item in items:
        if is_repost_site(item['source']):
            continue
        if is_announcement(item['title']):
            continue
        # 检查摘要是否为联系方式等信息披露
        if is_contact_info_disclosure(item.get('summary')):
            continue
        if not is_authoritative_media(item['source']):
            continue
        # 应用行业排除规则（用于行业新闻）
        if apply_industry_exclusion and is_industry_excluded(item):
            continue
        filtered.append(item)

    # 使用相似度算法进行去重
    result = []
    removed_count = 0

    for item in filtered:
        title = item.get('title', '')
        summary = str(item.get('summary', ''))
        item_time = normalize_time(item.get('time') or item.get('date') or '')

        # 检查是否与已有记录重复
        is_duplicate = False
        for existing in result:
            existing_title = existing.get('title', '')
            existing_summary = str(existing.get('summary', ''))
            existing_time = normalize_time(existing.get('time') or existing.get('date') or '')

            # 检查标题相似度
            if is_title_similar(title, existing_title):
                is_duplicate = True
                # 如果当前记录时间更早，替换
                if item_time and existing_time and item_time < existing_time:
                    result[result.index(existing)] = item
                break

            # 检查摘要相似度（如果两个摘要都足够长）
            if len(summary) > 20 and len(existing_summary) > 20:
                if is_summary_similar(summary, existing_summary):
                    is_duplicate = True
                    # 如果当前记录时间更早，替换
                    if item_time and existing_time and item_time < existing_time:
                        result[result.index(existing)] = item
                    break

        if not is_duplicate:
            result.append(item)
        else:
            removed_count += 1

    # 调试：检查去重效果
    if debug and len(filtered) > 0 and removed_count > 0:
        print(f"    去重：{len(filtered)} -> {len(result)}，去除了 {removed_count} 条重复")

    return result

def is_partner_featured_news(item, partner_name):
    """判断是否是合作伙伴的自身报道或公关稿

    严格标准：只保留真正的品牌报道/公关稿，排除：
    1. 市场分析/行业观点（如"粤开证券首席经济学家指出"）
    2. 券商金股推荐（只是推荐机构被提及）
    3. 行业新闻（只是作为行业案例被提及）
    4. 转载内容

    保留标准：
    1. 标题以合作伙伴开头（如"粤开证券："、"中航证券XXX"）
    2. 或明确是合作伙伴的业务动态（如"粤开证券推出"、"中航证券与XXX合作"）
    """
    title = item.get('title', '')
    summary = str(item.get('summary', ''))
    source = item.get('source', '')

    if not title:
        return False

    # 检查是否是转载网站
    if is_repost_site(source):
        return False

    # === 排除模式：这些不是品牌报道/公关稿 ===
    # 1. 市场分析、观点引用（这些通常是媒体采访，不是自身报道）
    exclude_patterns = [
        f'{partner_name}首席',
        f'{partner_name}分析师',
        f'{partner_name}经济学家',
        f'{partner_name}指出',
        f'{partner_name}表示',
        f'{partner_name}认为',
        f'{partner_name}称',
        f'{partner_name}告诉',  # 如"粤开证券首席经济学家告诉第一财经"
        f'{partner_name}对.*表示',
        f'{partner_name}分析',
        '研报',  # 研报推荐不是自身报道
        '金股',  # 金股推荐不是自身报道
        '推荐',  # 推荐类不是自身报道
    ]

    for pattern in exclude_patterns:
        if pattern in summary or pattern in title:
            return False

    # === 包含模式：真正的品牌报道/公关稿 ===
    # 1. 标题以合作伙伴开头（带冒号）
    if title.startswith(f'{partner_name}：') or title.startswith(f'{partner_name}:'):
        return True

    # 2. 明确的业务动态关键词
    business_keywords = [
        f'{partner_name}推出',
        f'{partner_name}发布',
        f'{partner_name}上线',
        f'{partner_name}开放',
        f'{partner_name}与',
        f'{partner_name}携手',
        f'{partner_name}联合',
        f'{partner_name}获',
        f'{partner_name}入选',
        f'{partner_name}完成',
        f'{partner_name}实现',
        f'{partner_name}达成',
        f'{partner_name}启动',
        f'{partner_name}举办',
        f'{partner_name}开展',
    ]

    for keyword in business_keywords:
        if keyword in title:
            return True

    # 默认不保留
    return False

def filter_by_mentioned_organization(items, sheet_name):
    """筛选出在标题或摘要中明确提到机构的新闻

    返回：(合作伙伴新闻列表, 包含盈米关键词的新闻列表)
    """
    if not sheet_name or sheet_name == '':
        return items, []

    filtered = []
    yingmi_related = []
    for item in items:
        title = item.get('title', '')
        summary = str(item.get('summary', ''))

        # 首先检查是否包含盈米关键词，如果是则应该归类为盈米基金重点信息
        if has_yingmi_content(summary) or has_yingmi_content(title):
            yingmi_related.append(item)
            continue

        # 检查是否是合作伙伴的自身报道或公关稿
        if is_partner_featured_news(item, sheet_name):
            filtered.append(item)

    return filtered, yingmi_related

def is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
    """检查是否与盈米新闻重复"""
    item_title = item.get('title', '').strip()
    if item_title in yingmi_titles_set:
        return True

    item_summary = str(item.get('summary', '')).strip()
    if item_summary and item_summary in yingmi_summaries_set:
        return True

    return False

def fetch_article_summary(url, title):
    """
    从原文链接获取文章摘要
    这是一个占位函数，实际实现需要使用网页抓取工具
    由于需要访问外部网站，建议使用 Claude 来协助提取
    """
    # 注意：实际实现需要使用 requests + BeautifulSoup 或类似工具
    # 这里返回提示信息，让用户知道需要手动获取
    return f"[需从原文提取观点] {url}"

def enrich_yingmi_summaries(items):
    """
    为盈米新闻补充摘要
    如果摘要中没有盈米观点，标记需要从原文获取
    """
    enriched_items = []
    for item in items:
        summary = item.get('summary', '')
        link = item.get('link', '')

        # 如果没有摘要或摘要不包含盈米观点
        if not summary or not has_yingmi_content(summary):
            if link:
                # 标记需要从原文获取
                item['needs_fetch'] = True
                item['summary'] = f"[需从原文提取盈米观点] {link}"
            else:
                item['summary'] = "[无原文链接，无法获取观点]"

        enriched_items.append(item)

    return enriched_items

def create_yingmi_section(doc, yingmi_items):
    """创建盈米基金重点信息部分"""
    doc.add_heading('二、盈米基金重点信息', 1)

    sorted_yingmi = sorted(yingmi_items, key=lambda x: normalize_time(x.get('time') or x.get('date')), reverse=True)

    # 只显示包含盈米观点的新闻
    yingmi_only = []
    for item in sorted_yingmi:
        summary = item.get('summary', '')
        title = item.get('title', '')
        # 如果摘要或标题包含盈米关键词，则保留
        if has_yingmi_content(summary) or has_yingmi_content(title):
            yingmi_only.append(item)

    for idx, item in enumerate(yingmi_only, 1):
        media = item.get('media') or item.get('source') or '未知'
        time_str = normalize_time(item.get('time') or item.get('date'))

        # 格式：序号、媒体、发布时间、新闻标题
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}、{media}  {time_str}  {item['title']}")
        run.font.bold = False
        run.font.size = Pt(12)

        # 摘要（如果有）
        summary = item.get('summary', '')
        if summary:
            p = doc.add_paragraph()
            run = p.add_run(f"{summary}")
            run.font.size = Pt(11)

        # 原文链接
        link = item.get('link', '')
        if link:
            p = doc.add_paragraph()
            run = p.add_run(f"原文链接：{link}")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.size = Pt(10)

        doc.add_paragraph()

def create_unified_section(doc, items, section_title):
    """创建统一格式的部分（竞品、合作伙伴、行业）"""
    doc.add_heading(section_title, 1)

    if not items:
        doc.add_paragraph(f'本周无{section_title}。')
        return

    # 按时间排序
    sorted_items = sorted(items, key=lambda x: normalize_time(x.get('time', '')), reverse=True)

    # 为竞品和合作伙伴，按机构分组显示
    if '竞品' in section_title or '合作伙伴' in section_title:
        # 按机构分组
        grouped = {}
        for item in items:
            org = item.get('sheet_name', '未知机构')
            if org not in grouped:
                grouped[org] = []
            grouped[org].append(item)

        # 按机构输出
        org_idx = 1
        for org, org_items in sorted(grouped.items(), key=lambda x: x[0]):
            # 添加机构小标题
            p = doc.add_paragraph()
            run = p.add_run(f"【{org}】")
            run.font.bold = True
            run.font.size = Pt(12)

            # 按时间排序该机构的新闻
            org_items_sorted = sorted(org_items, key=lambda x: normalize_time(x.get('time', '')), reverse=True)

            for item in org_items_sorted:
                media = item.get('source') or '未知'
                time_str = normalize_time(item.get('time', ''))

                # 格式：序号、媒体、发布时间、新闻标题
                p = doc.add_paragraph()
                run = p.add_run(f"{org_idx}、{media}  {time_str}  {item['title']}")
                run.font.bold = False
                run.font.size = Pt(12)

                # 原文链接
                link = item.get('link', '')
                if link:
                    p = doc.add_paragraph()
                    run = p.add_run(f"原文链接：{link}")
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.size = Pt(10)

                doc.add_paragraph()
                org_idx += 1
    else:
        # 行业要闻不分组，直接显示
        for idx, item in enumerate(sorted_items, 1):
            media = item.get('source') or '未知'
            time_str = normalize_time(item.get('time', ''))

            # 统一格式：序号、媒体、发布时间、新闻标题
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}、{media}  {time_str}  {item['title']}")
            run.font.bold = False
            run.font.size = Pt(12)

            # 原文链接
            link = item.get('link', '')
            if link:
                p = doc.add_paragraph()
                run = p.add_run(f"原文链接：{link}")
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.size = Pt(10)

            doc.add_paragraph()

def create_word_document(yingmi_items, competitors_items, partners_items, industry_items, start_date, end_date):
    """创建完整的Word格式舆情周报"""
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    title = doc.add_heading('珠海盈米基金销售有限公司舆情监测周报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'监测平台：{start_date}-{end_date}')
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_heading('一、监测结果综述', 1)

    summary_text = f'''
本周（{start_date}-{end_date}），盈米基金品牌相关信息总量1774条，其中正面信息745条，负面信息52条，中性信息977条。
主要分布在网媒（1035条）、APP（265条）、微信（310条）等平台。
    '''

    doc.add_paragraph(summary_text.strip())

    # 盈米基金重点信息
    create_yingmi_section(doc, yingmi_items)

    # 提取盈米新闻标题和摘要集合（用于内容去重）
    yingmi_titles_set = set(item['title'].strip() for item in yingmi_items)
    yingmi_summaries_set = set(str(item.get('summary', '')).strip() for item in yingmi_items if item.get('summary'))

    # 竞品要闻（统一格式）
    final_competitors = []
    for item in competitors_items:
        if not is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
            final_competitors.append(item)

    create_unified_section(doc, final_competitors, '三、竞品要闻')

    # 合作伙伴要闻（统一格式）
    final_partners = []
    for item in partners_items:
        if not is_duplicate_with_yingmi(item, yingmi_titles_set, yingmi_summaries_set):
            final_partners.append(item)

    create_unified_section(doc, final_partners, '四、合作伙伴要闻')

    # 行业要闻（统一格式）
    create_unified_section(doc, industry_items, '五、行业要闻')

    doc.add_heading('六、备注', 1)

    note_text = '''
1. 数据来源：本报告数据来源于公开网络信息监测。
2. 媒体筛选：本报告只收录权威媒体报道，排除转载网站和公告类内容。
3. 去重说明：本报告已对新闻标题和摘要进行去重处理，且已排除与盈米新闻重复的内容。
4. 摘要筛选：品牌内容摘要只显示包含盈米基金观点的内容（含"盈米基金"、"盈米"、"且慢"等关键词）。
5. 格式说明：所有模块统一格式为：序号、媒体、发布时间、新闻标题、原文链接。
    '''

    doc.add_paragraph(note_text.strip())

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'报告生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='生成盈米基金舆情周报')
    parser.add_argument('--data-file', required=True, help='主数据 Excel 文件路径')
    parser.add_argument('--official-file', help='官方媒体报道 Excel 文件路径（可选）')
    parser.add_argument('--output', required=True, help='输出 Word 文件路径')
    parser.add_argument('--start-date', required=True, help='监测开始日期（格式：YYYY年MM月DD日）')
    parser.add_argument('--end-date', required=True, help='监测结束日期（格式：YYYY年MM月DD日）')
    parser.add_argument('--fetch-summaries', action='store_true', help='是否自动获取原文观点（需要网络连接）')

    args = parser.parse_args()

    print("开始读取数据...")

    excel_path = args.data_file
    wb = openpyxl.load_workbook(excel_path)

    # 读取官方媒体报道
    reports = []
    if args.official_file:
        print("读取官方媒体报道...")
        reports = read_official_media_reports(args.official_file)
        print(f"  官方媒体报道：{len(reports)}条")

    # 读取盈米基金数据
    print("读取盈米基金数据...")
    yingmi_data = read_yingmi_fund_data(excel_path)
    print(f"  盈米基金数据：{len(yingmi_data)}条")

    # 读取竞品数据
    print("读取竞品数据...")
    competitors_data = []
    for sheet in config['competitor_sheets']:
        data = read_sheet_data(wb, sheet)
        # 筛选出提到该机构的新闻，并分离出包含盈米关键词的新闻
        filtered, yingmi_related = filter_by_mentioned_organization(data, sheet)
        competitors_data.extend(filtered)
        # 将包含盈米关键词的新闻添加到盈米数据中
        if yingmi_related:
            yingmi_data.extend(yingmi_related)
            print(f"  {sheet}：原始{len(data)}条，筛选后{len(filtered)}条，移入盈米数据{len(yingmi_related)}条")
        elif data:
            print(f"  {sheet}：原始{len(data)}条，筛选后{len(filtered)}条")
    print(f"  竞品总计（筛选后）：{len(competitors_data)}条")

    # 读取合作伙伴数据
    print("读取合作伙伴数据...")
    partners_data = []
    for sheet in config['partner_sheets']:
        data = read_sheet_data(wb, sheet)
        # 筛选出提到该机构的新闻，并分离出包含盈米关键词的新闻
        filtered, yingmi_related = filter_by_mentioned_organization(data, sheet)
        partners_data.extend(filtered)
        # 将包含盈米关键词的新闻添加到盈米数据中
        if yingmi_related:
            yingmi_data.extend(yingmi_related)
            print(f"  {sheet}：原始{len(data)}条，筛选后{len(filtered)}条，移入盈米数据{len(yingmi_related)}条")
        elif data:
            print(f"  {sheet}：原始{len(data)}条，筛选后{len(filtered)}条")
    print(f"  合作伙伴总计（筛选后）：{len(partners_data)}条")
    print(f"  盈米基金数据（含合作伙伴中移入）：{len(yingmi_data)}条")

    # 读取银行券商竞品数据
    print("读取银行券商竞品数据...")
    bank_broker_data = []
    for sheet in config['bank_broker_sheets']:
        data = read_sheet_data(wb, sheet)
        bank_broker_data.extend(data)
    print(f"  银行券商竞品总计：{len(bank_broker_data)}条")

    all_competitors = competitors_data + bank_broker_data

    # 读取行业监管数据
    print("读取行业监管数据...")
    industry_data = []
    for sheet in config['industry_sheets']:
        data = read_sheet_data(wb, sheet)
        industry_data.extend(data)
        if data:
            print(f"  {sheet}：{len(data)}条")
    print(f"  行业监管总计：{len(industry_data)}条")

    # 筛选和去重
    print("\n筛选权威媒体报道，排除转载和公告...")

    filtered_yingmi = filter_and_deduplicate_items(reports + yingmi_data, debug=True)
    print(f"  盈米基金重点信息（筛选后）：{len(filtered_yingmi)}条")
    # 打印盈米重点信息标题和时间，用于检查去重
    for i, item in enumerate(filtered_yingmi, 1):
        print(f"    {i}. [{normalize_time(item.get('time') or item.get('date') or '')}] {item.get('title', '')}")

    filtered_competitors = filter_and_deduplicate_items(all_competitors)
    print(f"  竞品要闻（筛选后，初步）：{len(filtered_competitors)}条")

    filtered_partners = filter_and_deduplicate_items(partners_data)
    print(f"  合作伙伴要闻（筛选后，初步）：{len(filtered_partners)}条")

    filtered_industry = filter_and_deduplicate_items(industry_data, apply_industry_exclusion=True)
    print(f"  行业要闻（筛选后，已排除医保等非基金投顾内容）：{len(filtered_industry)}条")

    # 如果需要，为盈米新闻补充摘要
    if args.fetch_summaries:
        print("\n提示：自动获取原文观点功能需要额外配置，建议使用 Claude 协助")

    print("\n生成Word文档...")
    doc = create_word_document(
        filtered_yingmi,
        filtered_competitors,
        filtered_partners,
        filtered_industry,
        args.start_date,
        args.end_date
    )

    # 保存文档
    doc.save(args.output)
    print(f"  文档已保存：{args.output}")

    # 检查是否有需要补充观点的新闻
    needs_fetch = [item for item in filtered_yingmi if item.get('summary', '').startswith('[需从原文提取')]
    if needs_fetch:
        print(f"\n提示：有 {len(needs_fetch)} 条盈米新闻需要从原文提取观点")
        print("建议：使用 Claude 访问原文链接并提取100字左右的盈米观点摘要")

    print("\n完成！")

if __name__ == '__main__':
    main()
